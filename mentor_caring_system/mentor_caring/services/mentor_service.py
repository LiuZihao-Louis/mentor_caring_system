from __future__ import annotations

import os
from typing import Any, Optional

from docx import Document

from mentor_caring.models import Appointment, ExportResult, Record
from mentor_caring.services.base import BaseService
from mentor_caring.services.message_service import MessageService
from mentor_caring.services.utils import iter_30_minute_slots, require_non_blank


class MentorService(BaseService):
    def _mentor_group_ids(self, mentor_id: str) -> list[str]:
        return [g.group_id for g in self.store.groups.all() if g.mentor_id == mentor_id]

    def _assert_own_student(self, mentor_id: str, student_id: str):
        mentor_id = require_non_blank(mentor_id, "mentor_id")
        student_id = require_non_blank(student_id, "student_id")
        self.store.mentors.require(mentor_id, "mentor")
        student = self.store.students.require(student_id, "student")
        if student.group_id not in self._mentor_group_ids(mentor_id):
            raise PermissionError("mentor can only access students in own groups")
        return student

    def search_student_by_id(self, mentor_id: str, student_id: str) -> dict:
        student = self._assert_own_student(mentor_id, student_id)
        self.log(mentor_id, "search student information", f"Searched {student_id}", student_id=student_id)
        return student.to_dict()

    def create_interview_record(
        self,
        mentor_id: str,
        student_id: str,
        date: str,
        time: str,
        problem_statement: str,
        interview_summary: str,
        follow_up_action: str | None = None,
    ) -> Record:
        student = self._assert_own_student(mentor_id, student_id)
        date = require_non_blank(date, "date")
        time = require_non_blank(time, "time")
        problem_statement = require_non_blank(problem_statement, "problem_statement")
        interview_summary = require_non_blank(interview_summary, "interview_summary")
        record = Record(
            record_id=self.store.next_id("REC"),
            student_id=student_id,
            mentor_id=mentor_id,
            group_id=student.group_id or "",
            date=date,
            time=time,
            problem_statement=problem_statement,
            interview_summary=interview_summary,
            follow_up_action=(follow_up_action or "").strip() or "None",
        )
        self.store.records.add(record.record_id, record)
        self.log(mentor_id, "create interview record", f"Created {record.record_id}", student_id=student_id)
        return record

    def update_interview_record(self, mentor_id: str, record_id: str, fields: dict[str, Any]) -> Record:
        mentor_id = require_non_blank(mentor_id, "mentor_id")
        record_id = require_non_blank(record_id, "record_id")
        record = self.store.records.require(record_id, "record")
        self._assert_own_student(mentor_id, record.student_id)
        record.update(**fields)
        self.log(mentor_id, "update interview record", f"Updated {record_id}", student_id=record.student_id)
        return record

    def delete_interview_record(self, mentor_id: str, record_id: str) -> dict:
        mentor_id = require_non_blank(mentor_id, "mentor_id")
        record_id = require_non_blank(record_id, "record_id")
        record = self.store.records.require(record_id, "record")
        self._assert_own_student(mentor_id, record.student_id)
        self.store.records.remove(record_id)
        self.log(mentor_id, "delete interview record", f"Deleted {record_id}", student_id=record.student_id)
        return {"deleted": record_id}

    def create_available_slots(self, mentor_id: str, student_id: str, date: str, start_time: str, end_time: str) -> list[Appointment]:
        student = self._assert_own_student(mentor_id, student_id)
        date = require_non_blank(date, "date")
        appointments = []
        for s, e in iter_30_minute_slots(start_time, end_time):
            appt = Appointment(
                appointment_id=self.store.next_id("APT"),
                mentor_id=mentor_id,
                student_id=student_id,
                group_id=student.group_id or "",
                date=date,
                start_time=s,
                end_time=e,
            )
            self.store.appointments.add(appt.appointment_id, appt)
            appointments.append(appt)
        self.log(mentor_id, "create appointment", f"Created {len(appointments)} slots", student_id=student_id)
        return appointments

    def confirm_appointment(self, mentor_id: str, appointment_id: str, venue: str) -> Appointment:
        mentor_id = require_non_blank(mentor_id, "mentor_id")
        appointment_id = require_non_blank(appointment_id, "appointment_id")
        venue = require_non_blank(venue, "venue")
        appt = self.store.appointments.require(appointment_id, "appointment")
        if appt.mentor_id != mentor_id:
            raise PermissionError("only appointment mentor can confirm this appointment")
        if appt.status != "booked":
            raise ValueError("appointment must be booked before confirmation")
        appt.update_status("confirmed", venue)
        self.notify(mentor_id, f"Appointment {appointment_id} confirmed", "appointment", appointment_id)
        self.notify(appt.student_id, f"Appointment {appointment_id} confirmed at {venue}", "appointment", appointment_id)
        self.log(mentor_id, "confirm appointment", f"Confirmed {appointment_id}", student_id=appt.student_id)
        return appt

    def forward_special_case_to_coordinator(self, mentor_id: str, student_id: str, coordinator_id: str, description: str):
        self._assert_own_student(mentor_id, student_id)
        coordinator_id = require_non_blank(coordinator_id, "coordinator_id")
        coordinator = self.store.coordinators.require(coordinator_id, "coordinator")
        description = require_non_blank(description, "description")
        return MessageService(self.store).send_message(
            sender_id=mentor_id,
            receiver_ids=[coordinator.coordinator_id],
            content=description,
            message_type="special_case",
            related_student_id=student_id,
        )

    def export_interview_records(self, mentor_id: str, student_ids: list[str] | None = None) -> ExportResult:
        mentor_id = require_non_blank(mentor_id, "mentor_id")
        self.store.mentors.require(mentor_id, "mentor")
        group_student_ids = set()
        for group in self.store.groups.all():
            if group.mentor_id == mentor_id:
                group_student_ids.update(group.student_ids)
        if student_ids:
            for sid in student_ids:
                if sid not in group_student_ids:
                    raise PermissionError("selected student is not in mentor's groups")
            selected = set(student_ids)
        else:
            selected = group_student_ids
        records = [r for r in self.store.records.all() if r.student_id in selected]
        export_dir = os.path.abspath("exports")
        os.makedirs(export_dir, exist_ok=True)
        file_name = f"mentor_{mentor_id}_records.docx"
        path = os.path.join(export_dir, file_name)
        doc = Document()
        doc.add_heading("Mentor Caring Interview Records", 0)
        for record in records:
            student = self.store.students.get(record.student_id)
            doc.add_paragraph(f"Student Name: {student.name if student else record.student_id}")
            doc.add_paragraph(f"Interview record: {record.date} {record.time}")
            doc.add_paragraph("Problem statements:")
            doc.add_paragraph(record.problem_statement)
            doc.add_paragraph("Interview summary:")
            doc.add_paragraph(record.interview_summary)
            doc.add_paragraph("Follow-up actions:")
            doc.add_paragraph(record.follow_up_action)
            doc.add_paragraph("")
        doc.save(path)
        result = ExportResult(file_path=path, file_name=file_name)
        self.store.exports.add(file_name, result)
        self.log(mentor_id, "export Word file", f"Exported {file_name}")
        return result
