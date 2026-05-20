from __future__ import annotations

import os
from collections import defaultdict
from typing import Any, Optional

from docx import Document

from mentor_caring.models import (
    Department,
    ExportResult,
    Faculty,
    FacultyConsultant,
    Major,
    MCPGroup,
    MCPCoordinator,
    Mentor,
    Student,
    User,
    make_password,
    ROLE_COORDINATOR,
    ROLE_MENTOR,
    ROLE_STUDENT,
)
from mentor_caring.services.base import BaseService
from mentor_caring.services.excel_utils import read_xlsx_as_dicts
from mentor_caring.services.utils import require_non_blank


class FacultyConsultantService(BaseService):
    def _consultant(self, consultant_id: str):
        consultant_id = require_non_blank(consultant_id, "consultant_id")
        return self.store.faculty_consultants.require(consultant_id, "faculty consultant")

    def _assert_own_faculty(self, consultant_id: str, faculty_name: str):
        consultant = self._consultant(consultant_id)
        if consultant.faculty_name != faculty_name:
            raise PermissionError("consultant can manage only own faculty")
        return consultant

    def _major(self, major_name: str):
        for major in self.store.majors.all():
            if major.name == major_name:
                return major
        raise LookupError(f"major not found: {major_name}")

    def import_organization_units(self, rows: list[dict] | None = None, xlsx_path: str | None = None, consultant_id: str | None = None) -> dict:
        consultant = self._consultant(consultant_id) if consultant_id else None
        if xlsx_path:
            rows = read_xlsx_as_dicts(xlsx_path)
        rows = rows or []
        summary = {"created_faculties": 0, "created_departments": 0, "created_majors": 0, "skipped_duplicates": 0}
        for row in rows:
            faculty_name = require_non_blank(row.get("Faculty"), "Faculty")
            dept_name = require_non_blank(row.get("Department"), "Department")
            major_name = require_non_blank(row.get("Major"), "Major")
            if consultant and consultant.faculty_name != faculty_name:
                raise PermissionError("consultant can import only own faculty organization")
            if not self.store.faculties.exists(faculty_name):
                self.store.faculties.add(faculty_name, Faculty(name=faculty_name))
                summary["created_faculties"] += 1
            else:
                summary["skipped_duplicates"] += 1
            if not self.store.departments.exists(dept_name):
                self.store.departments.add(dept_name, Department(name=dept_name, faculty_name=faculty_name))
                self.store.faculties.get(faculty_name).departments.append(dept_name)
                summary["created_departments"] += 1
            else:
                summary["skipped_duplicates"] += 1
            if not self.store.majors.exists(major_name):
                self.store.majors.add(major_name, Major(name=major_name, department_name=dept_name, faculty_name=faculty_name))
                dept = self.store.departments.get(dept_name)
                if major_name not in dept.majors:
                    dept.majors.append(major_name)
                summary["created_majors"] += 1
            else:
                summary["skipped_duplicates"] += 1
        return summary

    def import_students_and_mentors(
        self,
        consultant_id: str,
        rows: list[dict] | None = None,
        academic_year: str = "2024-2025",
        xlsx_path: str | None = None,
    ) -> dict:
        consultant = self._consultant(consultant_id)
        if xlsx_path:
            rows = read_xlsx_as_dicts(xlsx_path)
        rows = rows or []
        summary = {"students_created": 0, "students_updated": 0, "mentors_created": 0, "groups_created": 0}
        for row in rows:
            sid = require_non_blank(row.get("Student ID"), "Student ID")
            sname = require_non_blank(row.get("Student Name"), "Student Name")
            major_name = require_non_blank(row.get("Major"), "Major")
            status = str(row.get("Status") or "normal").lower()
            group_id = require_non_blank(row.get("Group ID"), "Group ID")
            mentor_name = require_non_blank(row.get("Mentor"), "Mentor")
            office = str(row.get("Office") or "")
            mentor_email = require_non_blank(row.get("Mentor Email"), "Mentor Email")
            major = self._major(major_name)
            if major.faculty_name != consultant.faculty_name:
                raise PermissionError("consultant can import only own faculty data")
            mentor = self.store.find_mentor_by_email(mentor_email)
            if mentor is not None and mentor.faculty_name != consultant.faculty_name:
                raise PermissionError("mentor faculty mismatch")
            if mentor is None:
                mentor_id = self.store.next_id("MTR")
                mentor = Mentor(
                    mentor_id=mentor_id,
                    name=mentor_name,
                    email=mentor_email,
                    office=office,
                    faculty_name=major.faculty_name,
                    department_name=major.department_name,
                    user_id=mentor_id,
                )
                self.store.mentors.add(mentor_id, mentor)
                self.store.users.add(mentor_id, User(
                    id=mentor_id, name=mentor_name, account=mentor_email, password_hash=make_password("mentor123"),
                    role=ROLE_MENTOR, email=mentor_email, is_staff=True,
                    faculty_name=major.faculty_name, department_name=major.department_name,
                ))
                summary["mentors_created"] += 1
            if not self.store.groups.exists(group_id):
                self.store.groups.add(group_id, MCPGroup(
                    group_id=group_id,
                    academic_year=academic_year,
                    year_label=group_id.split("-")[-1] if "-" in group_id else "",
                    mentor_id=mentor.mentor_id,
                    major_name=major.name,
                    department_name=major.department_name,
                    faculty_name=major.faculty_name,
                    student_ids=[],
                ))
                summary["groups_created"] += 1
            else:
                group = self.store.groups.get(group_id)
                if group.faculty_name != consultant.faculty_name:
                    raise PermissionError("consultant can update only own faculty groups")
                group.mentor_id = mentor.mentor_id
            student_email = f"{sid}@mail.bnbu.edu.cn"
            if self.store.students.exists(sid):
                student = self.store.students.get(sid)
                old_group = self.store.groups.get(student.group_id)
                if old_group and old_group.group_id != group_id and sid in old_group.student_ids:
                    old_group.student_ids.remove(sid)
                student.name = sname
                student.email = student_email
                student.status = status
                student.group_id = group_id
                student.faculty_name = major.faculty_name
                student.department_name = major.department_name
                student.major_name = major.name
                summary["students_updated"] += 1
            else:
                student = Student(
                    student_id=sid, name=sname, email=student_email, status=status,
                    group_id=group_id, faculty_name=major.faculty_name,
                    department_name=major.department_name, major_name=major.name, user_id=sid,
                )
                self.store.students.add(sid, student)
                self.store.users.add(sid, User(
                    id=sid, name=sname, account=sid, password_hash=make_password("student123"),
                    role=ROLE_STUDENT, email=student_email, is_staff=False,
                    faculty_name=major.faculty_name, department_name=major.department_name,
                ))
                summary["students_created"] += 1
            group = self.store.groups.get(group_id)
            if sid not in group.student_ids:
                group.student_ids.append(sid)
        self.log(consultant_id, "import students and mentors", str(summary))
        return summary

    def change_mentor_of_group(self, consultant_id: str, group_id: str, new_mentor_id: str):
        consultant = self._consultant(consultant_id)
        group = self.store.groups.require(require_non_blank(group_id, "group_id"), "group")
        if group.faculty_name != consultant.faculty_name:
            raise PermissionError("consultant can manage only groups in own faculty")
        self.store.mentors.require(new_mentor_id, "mentor")
        new_mentor = self.store.mentors.get(new_mentor_id)
        if new_mentor.faculty_name != consultant.faculty_name:
            raise PermissionError("new mentor must belong to consultant faculty")
        group.mentor_id = new_mentor_id
        self.log(consultant_id, "change mentor", f"Group {group_id} mentor -> {new_mentor_id}")
        return group

    def designate_coordinator(self, consultant_id: str, department_name: str, coordinator_id: str):
        consultant = self._consultant(consultant_id)
        dept = self.store.departments.require(department_name, "department")
        if dept.faculty_name != consultant.faculty_name:
            raise PermissionError("consultant can manage only own faculty department")
        self.store.coordinators.require(coordinator_id, "coordinator")
        dept.coordinator_id = coordinator_id
        return dept

    def import_coordinators(self, consultant_id: str, rows: list[dict] | None = None, xlsx_path: str | None = None):
        consultant = self._consultant(consultant_id)
        if xlsx_path:
            rows = read_xlsx_as_dicts(xlsx_path)
        rows = rows or []
        summary = {"coordinators_created": 0, "departments_assigned": 0}
        for row in rows:
            name = require_non_blank(row.get("Coordinator Name"), "Coordinator Name")
            email = require_non_blank(row.get("Email"), "Email")
            dept_name = require_non_blank(row.get("Department"), "Department")
            dept = self.store.departments.require(dept_name, "department")
            if dept.faculty_name != consultant.faculty_name:
                raise PermissionError("consultant can import only own faculty coordinators")
            existing_user = self.store.find_user_by_email(email)
            coord_id = existing_user.id if existing_user else self.store.next_id("COR")
            if not self.store.coordinators.exists(coord_id):
                coord = MCPCoordinator(coordinator_id=coord_id, name=name, email=email, department_name=dept_name, user_id=coord_id)
                self.store.coordinators.add(coord_id, coord)
                self.store.users.add(coord_id, User(
                    id=coord_id, name=name, account=email, password_hash=make_password("coord123"),
                    role=ROLE_COORDINATOR, email=email, is_staff=True,
                    faculty_name=consultant.faculty_name, department_name=dept_name,
                ))
                summary["coordinators_created"] += 1
            dept.coordinator_id = coord_id
            summary["departments_assigned"] += 1
        return summary

    def add_student_to_group(self, consultant_id: str, student_id: str, group_id: str):
        consultant = self._consultant(consultant_id)
        student = self.store.students.require(student_id, "student")
        group = self.store.groups.require(group_id, "group")
        if student.faculty_name != consultant.faculty_name or group.faculty_name != consultant.faculty_name:
            raise PermissionError("consultant can manage only own faculty")
        old_group = self.store.groups.get(student.group_id)
        if old_group and old_group.group_id != group_id and student_id in old_group.student_ids:
            old_group.student_ids.remove(student_id)
        if student_id not in group.student_ids:
            group.student_ids.append(student_id)
        student.group_id = group_id
        return group

    def remove_student_from_group(self, consultant_id: str, student_id: str, group_id: str):
        consultant = self._consultant(consultant_id)
        student = self.store.students.require(student_id, "student")
        group = self.store.groups.require(group_id, "group")
        if student.faculty_name != consultant.faculty_name or group.faculty_name != consultant.faculty_name:
            raise PermissionError("consultant can manage only own faculty")
        if student_id in group.student_ids:
            group.student_ids.remove(student_id)
        if student.group_id == group_id:
            student.group_id = None
        return group

    def search_student_information(self, consultant_id: str, student_id: str, academic_year: str | None = None, mentor_name: str | None = None):
        consultant = self._consultant(consultant_id)
        student = self.store.students.require(student_id, "student")
        if student.faculty_name != consultant.faculty_name:
            raise PermissionError("consultant can search only own faculty students")
        records = [r for r in self.store.records.all() if r.student_id == student_id]
        if academic_year:
            records = [r for r in records if (self.store.groups.get(r.group_id) and self.store.groups.get(r.group_id).academic_year == academic_year)]
        if mentor_name:
            records = [r for r in records if (self.store.mentors.get(r.mentor_id) and self.store.mentors.get(r.mentor_id).name == mentor_name)]
        self.log(consultant_id, "search student information", f"Searched {student_id}", student_id=student_id)
        data = student.to_dict()
        data["records"] = [r.to_dict() for r in records]
        return data

    def search_mentor_information(self, consultant_id: str, keyword: str | None = None, email: str | None = None, group_id: str | None = None):
        consultant = self._consultant(consultant_id)
        mentors = [m for m in self.store.mentors.all() if m.faculty_name == consultant.faculty_name]
        if keyword:
            mentors = [m for m in mentors if keyword.lower() in m.name.lower()]
        if email:
            mentors = [m for m in mentors if email.lower() in m.email.lower()]
        results = []
        for mentor in mentors:
            groups = [g for g in self.store.groups.all() if g.mentor_id == mentor.mentor_id]
            if group_id:
                groups = [g for g in groups if g.group_id == group_id]
            if groups or not group_id:
                results.append({
                    "mentor": mentor.to_dict(),
                    "groups": [{"group_id": g.group_id, "student_ids": g.student_ids} for g in groups],
                })
        self.log(consultant_id, "search mentor information", f"Found {len(results)} mentors")
        return results

    def view_student_logs(self, consultant_id: str, student_id: str | None = None):
        consultant = self._consultant(consultant_id)
        logs = []
        for log in self.store.logs.all():
            if log.student_id:
                student = self.store.students.get(log.student_id)
                if student and student.faculty_name == consultant.faculty_name:
                    if student_id is None or log.student_id == student_id:
                        logs.append(log)
        return logs

    def export_information_to_word(
        self,
        consultant_id: str,
        academic_years: list[str] | None = None,
        department_name: str | None = None,
        major_name: str | None = None,
        mentor_name: str | None = None,
        student_name: str | None = None,
    ) -> dict:
        consultant = self._consultant(consultant_id)
        students = [s for s in self.store.students.all() if s.faculty_name == consultant.faculty_name]
        if department_name:
            students = [s for s in students if s.department_name == department_name]
        if major_name:
            students = [s for s in students if s.major_name == major_name]
        if student_name:
            students = [s for s in students if student_name.lower() in s.name.lower()]
        by_major: dict[str, list] = defaultdict(list)
        for student in students:
            by_major[student.major_name].append(student)
        export_dir = os.path.abspath("exports")
        os.makedirs(export_dir, exist_ok=True)
        files = []
        for major, major_students in by_major.items():
            file_name = f"{consultant.faculty_name}_{major}_records.docx"
            path = os.path.join(export_dir, file_name)
            doc = Document()
            doc.add_heading("Mentor Caring Records", 0)
            for student in major_students:
                doc.add_paragraph(f"Faculty: {student.faculty_name}")
                doc.add_paragraph(f"Department: {student.department_name}        Major: {student.major_name}")
                doc.add_paragraph(f"Student Name: {student.name}")
                for record in [r for r in self.store.records.all() if r.student_id == student.student_id]:
                    group = self.store.groups.get(record.group_id)
                    if academic_years and group and group.academic_year not in academic_years:
                        continue
                    mentor = self.store.mentors.get(record.mentor_id)
                    if mentor_name and mentor and mentor.name != mentor_name:
                        continue
                    doc.add_paragraph(f"Interview record: {record.date}")
                    doc.add_paragraph("Problem statements:")
                    doc.add_paragraph(record.problem_statement)
                    doc.add_paragraph("Interview summary:")
                    doc.add_paragraph(record.interview_summary)
                    doc.add_paragraph("Follow-up actions:")
                    doc.add_paragraph(record.follow_up_action)
            doc.save(path)
            result = ExportResult(file_path=path, file_name=file_name)
            self.store.exports.add(file_name, result)
            files.append(result.to_dict())
        self.log(consultant_id, "export Word file", f"Exported {len(files)} files")
        return {"files": files, "count": len(files)}
